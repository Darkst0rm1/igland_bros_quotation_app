"""Waiving a charge.

A waiver is a third thing, and the tests here are mostly about keeping it
distinct from the other two:

* a **discount** changes what something costs;
* **deleting** the charge loses the fact that it applied;
* a **waiver** keeps the amount, keeps the row, and collects nothing.

That is why ``amount`` is never rewritten. Every assertion about un-waiving
below checks the exact original figure comes back, because the alternative
design — zeroing the amount and remembering it somewhere — is the one that
loses money to a rounding or a forgotten field.
"""
from __future__ import annotations

import datetime as dt
from decimal import Decimal as D

import pytest

from modules import pricing_snapshot, quotation_service, revision_service
from modules.authorization import PermissionDenied
from modules.calculation_engine import ChargeInput, compute_totals
from modules.catalogue_service import create_product, create_variant, set_price
from modules.constants import (
    ChargeType,
    ContainerSize,
    FreightMethod,
    PriceTierCode,
    QuotationStatus,
    RoleCode,
    WaiverStatus,
)
from modules.customer_service import create_customer
from modules.models import QuotationCharge, ShippingLine
from modules.quotation_service import QuotationError
from modules.validation import CustomerInput, PriceInput, ProductInput, VariantInput

JAN = dt.date(2026, 1, 1)
QUOTE_DAY = dt.date(2026, 8, 16)


@pytest.fixture
def admin(make_auth_user):
    return make_auth_user(RoleCode.SYS_ADMIN.value)


@pytest.fixture
def sales(make_auth_user):
    return make_auth_user(RoleCode.SALES.value)


@pytest.fixture
def carrier(session, seeded):
    line = ShippingLine(name="Test Line", is_active=True)
    session.add(line)
    session.commit()
    return line


@pytest.fixture
def quotation(session, admin):
    """One line of $20,000."""
    customer = create_customer(
        session, admin,
        CustomerInput(customer_number="CUST-0300", company_name="Waiver Ltd"),
    )
    session.flush()
    quote = quotation_service.create_draft(
        session, admin, customer.id, quote_date=QUOTE_DAY
    )
    product = create_product(
        session, admin,
        ProductInput(
            item_number="WB-12", name='12" White', size_label='12" White',
            flute="B", depth_in=D("2"),
        ),
    )
    session.flush()
    variant = create_variant(
        session, admin, product.id,
        VariantInput(
            variant_item_number="WB-12-A",
            board_quality="WT110 HPFL115 KM135", case_pack=50,
        ),
    )
    set_price(
        session, admin,
        PriceInput(
            product_variant_id=variant.id,
            price_tier_code=PriceTierCode.STANDARD.value,
            price_per_pack=D("10.00"), effective_from=JAN,
        ),
    )
    quotation_service.add_line(
        session, admin, quote,
        product_variant_id=variant.id,
        price_tier_code=PriceTierCode.STANDARD.value,
        quantity_packs=D("2000"),
    )
    session.commit()
    return quote


def _charge(session, admin, quotation, charge_type, rate, **kwargs):
    charge = quotation_service.add_charge(
        session, admin, quotation, charge_type=charge_type, rate=D(rate), **kwargs
    )
    session.commit()
    return charge


@pytest.fixture
def four_charges(session, admin, quotation):
    """One of every shape a charge comes in, including a future one."""
    return {
        "dies": _charge(session, admin, quotation, ChargeType.CUTTING_DIES, "400",
                        description="Cutting Dies"),
        "setup": _charge(session, admin, quotation, ChargeType.SETUP, "150",
                         description="Setup fee"),
        "handling": _charge(session, admin, quotation, ChargeType.BROKERAGE, "75",
                            description="Customs handling"),
        "internal": _charge(session, admin, quotation, ChargeType.OTHER, "60",
                            description="Internal only", is_customer_visible=False),
    }


# --------------------------------------------------------------------------- #
# The engine
# --------------------------------------------------------------------------- #

class TestTheEngine:
    def test_a_waived_charge_contributes_nothing(self):
        totals = compute_totals(
            [],
            charges=[
                ChargeInput(rate=D("400"), is_waived=True),
                ChargeInput(rate=D("8800")),
            ],
        )
        assert totals.charges_total == D("8800.00")
        assert totals.charges_waived == D("400.00")
        assert totals.grand_total == D("8800.00")

    def test_a_waived_charge_is_out_of_the_tax_base(self):
        totals = compute_totals(
            [],
            charges=[
                ChargeInput(rate=D("400"), is_waived=True, is_taxable=True),
                ChargeInput(rate=D("600"), is_taxable=True),
            ],
            tax_rate_pct=D("13"),
        )
        assert totals.taxable_base == D("600.00")
        assert totals.tax_amount == D("78.00")
        assert totals.grand_total == D("678.00")

    def test_a_waiver_is_not_a_discount(self):
        """A discount reduces the subtotal; a waiver does not touch it."""
        waived = compute_totals(
            [], charges=[ChargeInput(rate=D("400"), is_waived=True)]
        )
        discounted = compute_totals(
            [], charges=[ChargeInput(rate=D("400"))],
            quotation_discount_amount=D("400"),
        )
        assert waived.charges_total == D("0")
        assert waived.quotation_discount == D("0.00")
        assert discounted.charges_total == D("400.00")
        assert discounted.quotation_discount == D("0.00")   # nothing to discount

    def test_the_charge_keeps_its_own_amount(self):
        from modules.calculation_engine import charge_amount, payable_charge_amount

        charge = ChargeInput(quantity=D("2"), rate=D("200"), is_waived=True)
        assert charge_amount(charge) == D("400.00")
        assert payable_charge_amount(charge) == D("0")


# --------------------------------------------------------------------------- #
# The service
# --------------------------------------------------------------------------- #

class TestWaiving:
    def test_it_works_on_any_charge_type(self, session, admin, quotation, four_charges):
        """No list of waivable types exists, so a new one needs no change."""
        for key, charge in four_charges.items():
            quotation_service.waive_charge_directly(
                session, admin, quotation, charge.id, "goodwill"
            )
            session.commit()
            assert charge.is_waived, key
        assert quotation.charges_total == D("0.00")
        assert quotation.grand_total == quotation.subtotal == D("20000.00")

    def test_waiving_one_leaves_the_others_alone(
        self, session, admin, quotation, four_charges
    ):
        others = {
            k: (c.amount, c.is_waived)
            for k, c in four_charges.items() if k != "dies"
        }
        quotation_service.waive_charge_directly(
            session, admin, quotation, four_charges["dies"].id, "goodwill"
        )
        session.commit()

        assert four_charges["dies"].is_waived
        assert {
            k: (c.amount, c.is_waived)
            for k, c in four_charges.items() if k != "dies"
        } == others
        # 150 + 75 + 60 still billed
        assert quotation.charges_total == D("285.00")
        assert quotation.grand_total == D("20285.00")

    def test_the_amount_is_not_rewritten(self, session, admin, quotation, four_charges):
        dies = four_charges["dies"]
        quotation_service.waive_charge_directly(
            session, admin, quotation, dies.id, "goodwill"
        )
        session.commit()
        assert dies.amount == D("400.00"), "the original was overwritten"

    def test_unwaiving_returns_the_exact_amount(
        self, session, admin, quotation, four_charges
    ):
        dies = four_charges["dies"]
        before = quotation.grand_total

        quotation_service.waive_charge_directly(
            session, admin, quotation, dies.id, "goodwill"
        )
        session.commit()
        assert quotation.grand_total == before - D("400.00")

        quotation_service.remove_charge_waiver(
            session, admin, quotation, dies.id, "billing it after all"
        )
        session.commit()
        assert dies.is_waived is False
        assert dies.amount == D("400.00")
        assert quotation.grand_total == before

    def test_waiving_an_already_waived_charge_is_refused(
        self, session, admin, quotation, four_charges
    ):
        """Not merely idempotent — refused, so a second click cannot re-decide.

        The old boolean took a repeat silently. With a decision attached, a
        second waive would overwrite who decided and when, which is the part
        somebody asks about afterwards.
        """
        dies = four_charges["dies"]
        quotation_service.waive_charge_directly(
            session, admin, quotation, dies.id, "goodwill"
        )
        session.commit()
        decided_at = dies.waiver_decided_at

        with pytest.raises(QuotationError, match="already been waived"):
            quotation_service.waive_charge_directly(
                session, admin, quotation, dies.id, "goodwill again"
            )

        assert dies.waiver_decided_at == decided_at
        assert dies.waiver_reason == "goodwill"
        assert quotation.charges_total == D("285.00")

    def test_the_totals_are_recalculated_immediately(
        self, session, admin, quotation, four_charges
    ):
        """Stored, not merely computed on read: history and approval use these."""
        quotation_service.waive_charge_directly(
            session, admin, quotation, four_charges["setup"].id, "goodwill"
        )
        session.commit()
        snap = pricing_snapshot.base(quotation)
        assert quotation.charges_total == snap.charges_total == D("535.00")
        assert quotation.grand_total == snap.grand_total
        assert snap.charges_waived == D("150.00")

    def test_a_charge_from_another_quotation_is_refused(
        self, session, admin, quotation, four_charges
    ):
        other = quotation_service.create_draft(
            session, admin, quotation.customer_id, quote_date=QUOTE_DAY
        )
        session.flush()
        with pytest.raises(QuotationError, match="not part of this quotation"):
            quotation_service.waive_charge_directly(
                session, admin, other, four_charges["dies"].id, "goodwill"
            )
        assert four_charges["dies"].is_waived is False

    def test_a_charge_that_has_gone_is_refused(
        self, session, admin, quotation, four_charges
    ):
        dies_id = four_charges["dies"].id
        quotation_service.remove_charge(session, admin, quotation, dies_id)
        session.commit()
        with pytest.raises(QuotationError, match="not part of this quotation"):
            quotation_service.waive_charge_directly(
                session, admin, quotation, dies_id, "goodwill"
            )

    def test_sales_cannot_waive_on_someone_else_s_quotation(
        self, session, sales, quotation, four_charges
    ):
        with pytest.raises(PermissionDenied):
            quotation_service.waive_charge_directly(
                session, sales, quotation, four_charges["dies"].id, "goodwill"
            )
        assert four_charges["dies"].is_waived is False

    def test_it_is_audited_with_the_reason(
        self, session, admin, quotation, four_charges
    ):
        from sqlalchemy import select

        from modules.models import AuditLog

        quotation_service.waive_charge_directly(
            session, admin, quotation, four_charges["dies"].id, "Goodwill on a first order",
        )
        session.commit()

        rows = [
            r for r in session.execute(select(AuditLog)).scalars()
            if r.reason == "Goodwill on a first order"
        ]
        assert rows, "the waiver was not audited"
        assert rows[0].new_value_json["waiver_status"] == "APPROVED"
        assert rows[0].old_value_json["waiver_status"] == "NONE"
        assert str(rows[0].new_value_json["amount"]) == "400.00"
        # Who asked and who decided, both recorded — the same person here,
        # because a manager waiving directly is a one-step waiver.
        assert rows[0].new_value_json["requested_by_id"] == admin.id
        assert rows[0].new_value_json["decided_by_id"] == admin.id


class TestWaivingFreight:
    """The derived charge, which nothing else writes by hand."""

    @pytest.fixture
    def with_freight(self, session, admin, quotation, carrier):
        from modules import shipping_service

        shipping_service.add_container(
            session, admin, quotation, shipping_line_id=carrier.id,
            container_size=ContainerSize.FORTY_FT_HC,
            container_count=D("2"), freight_cost=D("4400"),
        )
        session.commit()
        return quotation

    def _freight(self, session, quotation):
        return session.query(QuotationCharge).filter_by(
            quotation_id=quotation.id, source="shipment"
        ).one()

    def test_freight_can_be_waived_like_anything_else(
        self, session, admin, with_freight
    ):
        freight = self._freight(session, with_freight)
        assert with_freight.grand_total == D("28800.00")

        quotation_service.waive_charge_directly(
            session, admin, with_freight, freight.id, "goodwill"
        )
        session.commit()

        assert freight.amount == D("8800.00")
        assert with_freight.charges_total == D("0.00")
        assert with_freight.grand_total == D("20000.00")

    def test_resyncing_the_shipment_does_not_un_waive_it(
        self, session, admin, with_freight, carrier
    ):
        """sync_freight rewrites the charge on every container change.

        It must leave the waiver alone, or editing a port would quietly start
        billing freight the customer was told they would not pay.
        """
        from modules import shipping_service

        freight = self._freight(session, with_freight)
        quotation_service.waive_charge_directly(
            session, admin, with_freight, freight.id, "goodwill"
        )
        session.commit()

        shipping_service.update_shipment(
            session, admin, with_freight, port_of_loading="Istanbul"
        )
        session.commit()

        assert self._freight(session, with_freight).is_waived
        assert with_freight.grand_total == D("20000.00")

    def test_changing_the_freight_amount_keeps_the_waiver(
        self, session, admin, with_freight
    ):
        from modules import shipping_service

        freight = self._freight(session, with_freight)
        quotation_service.waive_charge_directly(
            session, admin, with_freight, freight.id, "goodwill"
        )
        session.commit()

        container = shipping_service.get_shipment(session, with_freight.id).containers[0]
        shipping_service.update_container(
            session, admin, with_freight, container.id, freight_cost=D("5000")
        )
        session.commit()

        updated = self._freight(session, with_freight)
        assert updated.is_waived
        assert updated.amount == D("10000.00"), "the amount should still track"
        assert with_freight.grand_total == D("20000.00")


# --------------------------------------------------------------------------- #
# The approval workflow
# --------------------------------------------------------------------------- #

class TestTheApprovalWorkflow:
    """An employee asks; a manager decides; the charge is billed in between."""

    @pytest.fixture
    def manager(self, make_auth_user):
        return make_auth_user(RoleCode.SALES_MANAGER.value)

    @pytest.fixture
    def requested(self, session, sales, admin, quotation, four_charges):
        quotation.sales_user_id = sales.id
        session.flush()
        quotation_service.request_charge_waiver(
            session, sales, quotation, four_charges["dies"].id,
            "Customer is a first order and asked for the dies to be covered",
        )
        session.commit()
        return four_charges["dies"]

    # --- requesting -------------------------------------------------------- #

    def test_a_pending_waiver_is_still_billed(self, session, quotation, requested):
        """The load-bearing rule. Asking is not receiving."""
        assert requested.waiver_status is WaiverStatus.PENDING
        assert requested.waiver_pending
        assert requested.is_waived is False
        assert quotation.charges_total == D("685.00")   # all four still charged
        assert quotation.grand_total == D("20685.00")

    def test_the_requester_and_the_reason_are_stored(
        self, session, sales, requested
    ):
        assert requested.waiver_requested_by_id == sales.id
        assert requested.waiver_requested_at is not None
        assert "first order" in requested.waiver_reason
        assert requested.waiver_decided_by_id is None
        assert requested.waiver_decided_at is None

    def test_a_reason_is_required(self, session, sales, quotation, four_charges):
        for empty in ("", "   "):
            with pytest.raises(QuotationError, match="reason is required"):
                quotation_service.request_charge_waiver(
                    session, sales, quotation, four_charges["setup"].id, empty
                )
        assert four_charges["setup"].waiver_status is WaiverStatus.NONE

    def test_requesting_twice_is_refused(self, session, sales, quotation, requested):
        with pytest.raises(QuotationError, match="already awaiting"):
            quotation_service.request_charge_waiver(
                session, sales, quotation, requested.id, "asking again"
            )

    def test_a_role_without_the_permission_cannot_request(
        self, session, make_auth_user, quotation, four_charges
    ):
        pricing_admin = make_auth_user(RoleCode.PRICING_ADMIN.value)
        with pytest.raises(PermissionDenied):
            quotation_service.request_charge_waiver(
                session, pricing_admin, quotation, four_charges["dies"].id, "why not"
            )

    # --- deciding ---------------------------------------------------------- #

    def test_sales_cannot_approve_their_own_request(
        self, session, sales, quotation, requested
    ):
        """The whole point of the split: asking must not be deciding."""
        with pytest.raises(PermissionDenied):
            quotation_service.approve_charge_waiver(
                session, sales, quotation, requested.id
            )
        assert requested.is_waived is False
        assert quotation.charges_total == D("685.00")

    def test_a_manager_approving_takes_the_money_off(
        self, session, manager, quotation, requested
    ):
        quotation_service.approve_charge_waiver(
            session, manager, quotation, requested.id, "Agreed, first order"
        )
        session.commit()

        assert requested.waiver_status is WaiverStatus.APPROVED
        assert requested.is_waived
        assert requested.amount == D("400.00"), "the amount was rewritten"
        assert requested.waiver_decided_by_id == manager.id
        assert requested.waiver_decided_at is not None
        assert requested.waiver_decision_note == "Agreed, first order"
        # The requester is preserved: it was not the manager who asked.
        assert requested.waiver_requested_by_id != manager.id
        assert quotation.charges_total == D("285.00")
        assert quotation.grand_total == D("20285.00")

    def test_a_manager_rejecting_goes_on_billing_it(
        self, session, manager, quotation, requested
    ):
        quotation_service.reject_charge_waiver(
            session, manager, quotation, requested.id, "Margin is already thin"
        )
        session.commit()

        assert requested.waiver_status is WaiverStatus.REJECTED
        assert requested.is_waived is False
        assert quotation.charges_total == D("685.00")
        assert quotation.grand_total == D("20685.00")
        # The refusal is kept, not cleared.
        assert requested.waiver_decision_note == "Margin is already thin"
        assert "first order" in requested.waiver_reason

    def test_a_rejected_waiver_can_be_asked_for_again(
        self, session, sales, manager, quotation, requested
    ):
        quotation_service.reject_charge_waiver(
            session, manager, quotation, requested.id
        )
        session.commit()
        quotation_service.request_charge_waiver(
            session, sales, quotation, requested.id, "the customer pushed back"
        )
        session.commit()
        assert requested.waiver_status is WaiverStatus.PENDING

    def test_deciding_something_nobody_asked_for_is_refused(
        self, session, manager, quotation, four_charges
    ):
        for decide in (
            quotation_service.approve_charge_waiver,
            quotation_service.reject_charge_waiver,
        ):
            with pytest.raises(QuotationError, match="no waiver awaiting"):
                decide(session, manager, quotation, four_charges["setup"].id)

    # --- direct waiving ---------------------------------------------------- #

    def test_a_manager_may_waive_without_a_request(
        self, session, manager, quotation, four_charges
    ):
        charge = four_charges["setup"]
        quotation_service.waive_charge_directly(
            session, manager, quotation, charge.id, "Covered as a gesture"
        )
        session.commit()

        assert charge.is_waived
        # Recorded as a one-step waiver rather than a fabricated request.
        assert charge.waiver_requested_by_id == manager.id
        assert charge.waiver_decided_by_id == manager.id
        assert charge.waiver_requested_at == charge.waiver_decided_at

    def test_a_direct_waiver_needs_a_reason(
        self, session, manager, quotation, four_charges
    ):
        with pytest.raises(QuotationError, match="reason is required"):
            quotation_service.waive_charge_directly(
                session, manager, quotation, four_charges["setup"].id, "  "
            )

    def test_sales_cannot_waive_directly(
        self, session, sales, quotation, four_charges
    ):
        with pytest.raises(PermissionDenied):
            quotation_service.waive_charge_directly(
                session, sales, quotation, four_charges["dies"].id, "please"
            )

    # --- un-waiving -------------------------------------------------------- #

    def test_unwaiving_needs_the_approver(
        self, session, sales, manager, quotation, requested
    ):
        quotation_service.approve_charge_waiver(
            session, manager, quotation, requested.id
        )
        session.commit()

        with pytest.raises(PermissionDenied):
            quotation_service.remove_charge_waiver(
                session, sales, quotation, requested.id, "billing it after all"
            )
        assert requested.is_waived

        quotation_service.remove_charge_waiver(
            session, manager, quotation, requested.id, "billing it after all"
        )
        session.commit()
        assert requested.waiver_status is WaiverStatus.NONE
        assert requested.amount == D("400.00")
        assert quotation.charges_total == D("685.00")

    def test_unwaiving_something_not_waived_is_refused(
        self, session, manager, quotation, four_charges
    ):
        with pytest.raises(QuotationError, match="not waived"):
            quotation_service.remove_charge_waiver(
                session, manager, quotation, four_charges["setup"].id
            )

    # --- the queue --------------------------------------------------------- #

    def test_the_queue_lists_what_is_awaiting_a_decision(
        self, session, manager, quotation, requested
    ):
        rows = quotation_service.pending_waivers(session, manager)
        assert [c.id for c, _ in rows] == [requested.id]
        assert rows[0][1].id == quotation.id

    def test_the_queue_empties_once_decided(
        self, session, manager, quotation, requested
    ):
        quotation_service.approve_charge_waiver(
            session, manager, quotation, requested.id
        )
        session.commit()
        assert quotation_service.pending_waivers(session, manager) == []

    def test_sales_cannot_see_the_queue(self, session, sales):
        with pytest.raises(PermissionDenied):
            quotation_service.pending_waivers(session, sales)

    # --- an issued quotation ----------------------------------------------- #

    def test_nothing_may_be_waived_on_an_issued_quotation(
        self, session, manager, quotation, four_charges
    ):
        revision_service.issue(session, manager, quotation)
        session.commit()
        with pytest.raises(QuotationError, match="has been issued"):
            quotation_service.waive_charge_directly(
                session, manager, quotation, four_charges["dies"].id, "too late"
            )

    def test_a_waiver_can_be_decided_while_the_quotation_awaits_approval(
        self, session, sales, manager, quotation, requested
    ):
        """The deadlock this avoids.

        ``require_edit_quotation`` allows only DRAFT and REVISION_REQUIRED, so
        gating waivers on it would mean a request raised on a draft could never
        be decided once the quotation was submitted — which is exactly when the
        manager is looking at it.
        """
        quotation_service.change_status(
            session, manager, quotation, QuotationStatus.PENDING_APPROVAL
        )
        session.commit()

        # The quotation is genuinely no longer editable at this point, which is
        # what would have blocked the decision.
        from modules.authorization import can_edit_quotation

        assert not can_edit_quotation(manager, quotation)

        quotation_service.approve_charge_waiver(
            session, manager, quotation, requested.id
        )
        session.commit()
        assert requested.is_waived
        assert quotation.charges_total == D("285.00")


class TestPendingNeverReachesTheCustomer:
    """A concession asked for is not one given, and must not look like one."""

    @pytest.fixture
    def pending(self, session, sales, quotation, four_charges):
        quotation.sales_user_id = sales.id
        session.flush()
        quotation_service.request_charge_waiver(
            session, sales, quotation, four_charges["dies"].id, "asked for"
        )
        session.commit()
        return quotation

    def test_the_document_shows_it_as_an_ordinary_charge(self, session, pending):
        from modules import document_model

        rows = {
            t.label: t.amount
            for t in document_model.build_document(session, pending).totals
        }
        assert rows["Cutting Dies"] == "$400.00"
        assert not any("WAIVED" in label for label in rows)
        assert not any("PENDING" in label.upper() for label in rows)
        assert rows["Total (USD)"] == "$20,685.00"

    def test_the_pdf_says_nothing_about_it(self, session, pending):
        from modules import document_model, pdf_generator

        text = _pdf_text(
            pdf_generator.render(document_model.build_document(session, pending))
        )
        assert "WAIVED" not in text
        assert "PENDING" not in text.upper()

    def test_the_portal_shows_it_as_an_ordinary_charge(self, session, pending):
        from portal import pdf_model

        snapshot = pricing_snapshot.base(pending)
        rows = pdf_model._charge_rows(pending, snapshot.charges_customer_visible)
        assert not any("WAIVED" in r.label for r in rows)
        assert not any(r.label == "Total charges" for r in rows)


# --------------------------------------------------------------------------- #
# Revisions
# --------------------------------------------------------------------------- #

class TestRevisions:
    def test_a_revision_inherits_the_waiver(
        self, session, admin, quotation, four_charges
    ):
        """A concession already given does not reappear as a charge."""
        quotation_service.waive_charge_directly(
            session, admin, quotation, four_charges["dies"].id, "goodwill"
        )
        revision_service.issue(session, admin, quotation)
        session.commit()
        assert quotation.grand_total == D("20285.00")

        revised = revision_service.create_revision(
            session, admin, quotation, "customer added a size"
        )
        session.commit()

        waived = [c for c in revised.charges if c.is_waived]
        assert len(waived) == 1
        assert waived[0].amount == D("400.00")
        assert revised.grand_total == D("20285.00")

    def test_unwaiving_on_the_revision_leaves_the_issued_copy_alone(
        self, session, admin, quotation, four_charges
    ):
        quotation_service.waive_charge_directly(
            session, admin, quotation, four_charges["dies"].id, "goodwill"
        )
        revision_service.issue(session, admin, quotation)
        session.commit()
        issued_total = quotation.grand_total

        revised = revision_service.create_revision(
            session, admin, quotation, "charging for the dies after all"
        )
        session.flush()
        copy = next(c for c in revised.charges if c.is_waived)
        quotation_service.remove_charge_waiver(
            session, admin, revised, copy.id, "billing it after all"
        )
        session.commit()

        assert revised.grand_total == D("20685.00")
        assert quotation.grand_total == issued_total == D("20285.00")


# --------------------------------------------------------------------------- #
# What the customer sees
# --------------------------------------------------------------------------- #

def _pdf_text(pdf: bytes) -> str:
    from io import BytesIO

    from pypdf import PdfReader

    return "\n".join(p.extract_text() or "" for p in PdfReader(BytesIO(pdf)).pages)


class TestTheDocument:
    @pytest.fixture
    def waived_dies(self, session, admin, quotation, carrier):
        """Exactly the brief's example: dies waived, freight charged.

        Deliberately not built on ``four_charges`` — the worked example has two
        charges and the arithmetic in it is the point.
        """
        from modules import shipping_service

        quotation_service.add_charge(
            session, admin, quotation, charge_type=ChargeType.CUTTING_DIES,
            rate=D("400"), description="Cutting Dies",
        )
        shipping_service.add_container(
            session, admin, quotation, shipping_line_id=carrier.id,
            container_size=ContainerSize.FORTY_FT_HC,
            container_count=D("2"), freight_cost=D("4400"),
        )
        session.flush()
        dies = session.query(QuotationCharge).filter_by(
            quotation_id=quotation.id, charge_type=ChargeType.CUTTING_DIES
        ).one()
        quotation_service.waive_charge_directly(
            session, admin, quotation, dies.id, "goodwill"
        )
        session.commit()
        return quotation

    def test_the_row_stays_at_its_amount_and_is_marked(self, session, waived_dies):
        from modules import document_model

        rows = {
            t.label: t.amount
            for t in document_model.build_document(session, waived_dies).totals
        }
        assert rows["Cutting Dies — WAIVED"] == "$400.00"
        assert rows["Ocean freight — 2 containers"] == "$8,800.00"
        assert rows["Total charges"] == "$8,800.00"
        assert rows["Total (USD)"] == "$28,800.00"

    def test_the_worked_example_from_the_brief(self, session, waived_dies):
        """Cutting Dies $400 WAIVED, Ocean Freight $8,800, charges $8,800."""
        from modules import document_model, pdf_generator

        text = _pdf_text(
            pdf_generator.render(
                document_model.build_document(session, waived_dies)
            )
        )
        assert "WAIVED" in text
        assert "$400.00" in text
        assert "$8,800.00" in text
        assert "$28,800.00" in text

    def test_the_word_document_says_the_same(self, session, waived_dies):
        from io import BytesIO

        from docx import Document as DocxDocument

        from modules import document_model, docx_generator

        model = document_model.build_document(session, waived_dies)
        doc = DocxDocument(BytesIO(docx_generator.render(model)))
        body = " ".join(p.text for p in doc.paragraphs)
        cells = " ".join(
            c.text for t in doc.tables for r in t.rows for c in r.cells
        )
        assert "WAIVED" in body + cells

    def test_nothing_is_marked_when_nothing_is_waived(
        self, session, admin, quotation, four_charges
    ):
        from modules import document_model

        doc = document_model.build_document(session, quotation)
        assert not any("WAIVED" in t.label for t in doc.totals)
        assert not any(t.label == "Total charges" for t in doc.totals), (
            "the extra row is noise when the rows already add up"
        )

    def test_an_internal_charge_that_is_waived_is_not_counted_twice(
        self, session, admin, quotation, four_charges
    ):
        """It is invisible *and* waived: it must not reappear in the fold-in."""
        from modules import document_model

        quotation_service.waive_charge_directly(
            session, admin, quotation, four_charges["internal"].id, "goodwill"
        )
        session.commit()

        doc = document_model.build_document(session, quotation)
        assert not any(t.label == "Additional charges" for t in doc.totals)
        assert quotation.charges_total == D("625.00")   # 400 + 150 + 75


class TestThePortal:
    """The customer-facing PDF. The HTML projection is covered in
    ``test_portal_web``, where the token and send fixtures already live."""

    @pytest.fixture
    def waived(self, session, admin, quotation, four_charges):
        quotation_service.waive_charge_directly(
            session, admin, quotation, four_charges["dies"].id, "goodwill"
        )
        session.commit()
        return quotation

    def test_the_portal_pdf_marks_it(self, session, waived):
        from portal import pdf_model

        snapshot = pricing_snapshot.base(waived)
        rows = pdf_model._charge_rows(waived, snapshot.charges_customer_visible)
        labels = [r.label for r in rows]
        assert any("WAIVED" in label for label in labels)
        assert "Total charges" in labels
        by_label = {r.label: r.amount for r in rows}
        assert by_label["Cutting Dies — WAIVED"] == D("400.00")
        assert by_label["Total charges"] == D("225.00")
