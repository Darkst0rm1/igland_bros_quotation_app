"""Documents (PDF and Word), the release gate, approvals and revisions."""

from __future__ import annotations

import datetime as dt
from decimal import Decimal as D
from io import BytesIO

import pytest

from modules import (
    approval_service,
    document_model,
    document_service,
    docx_generator,
    pdf_generator,
    quotation_service,
    revision_service,
)
from modules.approval_service import ApprovalError
from modules.authorization import PermissionDenied
from modules.catalogue_service import create_product, create_variant, set_cost, set_price
from modules.constants import (
    AuditAction,
    ChargeType,
    PriceTierCode,
    QuotationStatus,
    RoleCode,
)
from modules.customer_service import add_address, add_contact, create_customer
from modules.document_service import DocumentFormat
from modules.models import AuditLog, ImmutableRecordError, Quotation
from modules.revision_service import RevisionError
from modules.validation import (
    AddressInput,
    ContactInput,
    CostInput,
    CustomerInput,
    PriceInput,
    ProductInput,
    VariantInput,
)

JAN = dt.date(2026, 1, 1)
QUOTE_DAY = dt.date(2026, 8, 3)


# --------------------------------------------------------------------------- #
# Fixtures
# --------------------------------------------------------------------------- #

@pytest.fixture
def sales(make_auth_user):
    return make_auth_user(RoleCode.SALES.value, username="alice")


@pytest.fixture
def manager(make_auth_user):
    return make_auth_user(RoleCode.SALES_MANAGER.value, username="mgr")


@pytest.fixture
def admin(make_auth_user):
    return make_auth_user(RoleCode.SYS_ADMIN.value, username="root")


@pytest.fixture
def variant(session, admin):
    product = create_product(
        session, admin,
        ProductInput(
            item_number="WB-12", name='12" White', size_label='12" White',
            flute="B", depth_in=D("2"),
        ),
    )
    session.flush()
    made = create_variant(
        session, admin, product.id,
        VariantInput(
            variant_item_number="WB-12-115",
            board_quality="WT110 HPFL115 KM135",
            case_pack=50,
        ),
    )
    for tier, price in (
        (PriceTierCode.STANDARD.value, "7.42"),
        (PriceTierCode.EIGHT_CONTAINER.value, "6.98"),
    ):
        set_price(
            session, admin,
            PriceInput(
                product_variant_id=made.id, price_tier_code=tier,
                price_per_pack=D(price), effective_from=JAN,
            ),
        )
    session.commit()
    return made


@pytest.fixture
def quotation(session, sales, admin, variant):
    customer = create_customer(
        session, sales,
        CustomerInput(customer_number="CUST-0001", company_name="Bunzl Canada"),
    )
    session.flush()
    add_contact(
        session, sales, customer.id,
        ContactInput(name="Michel Dupont", email="michel@bunzl.invalid", is_primary=True),
    )
    add_address(
        session, sales, customer.id,
        AddressInput(
            address_type="BILLING", line1="100 Bay Street", city="Toronto",
            country="Canada", is_default=True,
        ),
    )
    quote = quotation_service.create_draft(
        session, sales, customer.id,
        project_name="Pizza Box Program", quote_date=QUOTE_DAY,
    )
    quotation_service.add_line(
        session, sales, quote,
        product_variant_id=variant.id,
        price_tier_code=PriceTierCode.STANDARD.value,
        quantity_packs=D("1000"),
    )
    quotation_service.add_plate_charge(
        session, sales, quote, number_of_sizes=1, number_of_colours=4
    )
    session.commit()
    return quote


@pytest.fixture
def carrier(session, admin):
    from modules.models import ShippingLine

    line = ShippingLine(name="Maersk", is_active=True)
    session.add(line)
    session.commit()
    return line


def _approve_and_issue(session, quotation, sales, manager):
    """Push a draft through submission and approval."""
    approval = approval_service.submit(session, quotation, sales)
    session.commit()
    if approval is not None:
        approval_service.approve(session, quotation, manager, approval.id)
        session.commit()
    return quotation


# --------------------------------------------------------------------------- #
# The document model
# --------------------------------------------------------------------------- #

class TestDocumentModel:
    def test_it_carries_the_line_snapshot_not_the_live_catalogue(
        self, session, quotation
    ):
        model = document_model.build_document(session, quotation)
        assert len(model.lines) == 1
        values = model.lines[0].values
        assert values["size"] == '12" White'
        assert values["quantity_packs"] == "1,000"
        assert values["line_total"] == "$7,420.00"

    def test_the_board_quality_is_on_the_document_and_the_size_is_not_doubled(
        self, session, quotation
    ):
        """Two faults in one column, found by looking at a real quotation.

        ``description`` fell back to ``size_label`` whenever no override was
        typed — which is almost always — so the printed table carried the same
        value twice under two headings. The board quality, meanwhile, reached
        no default layout at all, and it is the part that identifies the
        product: WTL125 FL120 IK120 and IK135 are different goods at different
        prices on the same size.
        """
        columns = document_model.DEFAULT_COLUMNS
        assert "board_quality" in columns
        assert "description" not in columns
        assert columns.index("board_quality") < columns.index("size")

        model = document_model.build_document(session, quotation)
        values = model.lines[0].values
        assert values["board_quality"] == "WT110 HPFL115 KM135"
        assert values["size"] == '12" White'
        assert values["board_quality"] != values["size"]

    def test_a_typed_description_and_the_remarks_still_reach_the_document(
        self, session, sales, quotation
    ):
        """Dropping the description column must not orphan its two fields.

        Both are offered on the line editor. If neither printed anywhere the
        boxes would be controls that do nothing, which is the shape of bug this
        codebase keeps finding.
        """
        from modules import quotation_service

        line = quotation.items[0]
        quotation_service.update_line(
            session, sales, quotation, line.id,
            description_override="Printed one colour, food grade",
            customer_remarks="Delivered in two drops",
        )
        session.commit()

        values = document_model.build_document(session, quotation).lines[0].values
        assert "Printed one colour, food grade" in values["size"]
        assert "Delivered in two drops" in values["size"]
        # The board quality column is unaffected by either.
        assert values["board_quality"] == "WT110 HPFL115 KM135"

    def test_containers_are_derived_from_the_quantity_when_none_was_typed(
        self, session, quotation, variant
    ):
        """The customer is told what their quantity fills, not left guessing.

        Nobody types a container count on most lines, so the column would
        otherwise be permanently blank. The estimate is presentational: it is
        never written back to ``container_count``, because
        ``pricing_service._quotation_container_total`` trusts a typed count
        above a catalogue estimate when choosing the container price tier.
        """
        from decimal import Decimal as D

        from modules.constants import ContainerSize, ContainerType
        from modules.models import ProductContainerCapacity, Product

        line = quotation.items[0]
        assert not line.container_count, "fixture should have none typed"

        product = session.get(Product, variant.product_id)
        product.units_per_bundle = D("10")           # 10 packs to a bundle
        session.add(ProductContainerCapacity(
            product_id=product.id,
            container_size=ContainerSize.FORTY_FT_HC,
            container_type=ContainerType.DRY,
            bundles_per_container=D("20"),           # -> 200 packs per container
        ))
        session.flush()

        values = document_model.build_document(session, quotation).lines[0].values
        # 20 bundles x 10 pieces = 200 pieces per container, / 50 per pack
        # = 4 packs per container; 1,000 packs / 4 = 250.
        assert values["containers"] == "250"
        assert not line.container_count, "the estimate must not be written back"

    def test_a_typed_container_count_beats_the_estimate(
        self, session, quotation, variant
    ):
        """Somebody's own statement of the shipment outranks the workbook."""
        from decimal import Decimal as D

        from modules.constants import ContainerSize, ContainerType
        from modules.models import ProductContainerCapacity, Product

        product = session.get(Product, variant.product_id)
        product.units_per_bundle = D("10")
        session.add(ProductContainerCapacity(
            product_id=product.id,
            container_size=ContainerSize.FORTY_FT_HC,
            container_type=ContainerType.DRY,
            bundles_per_container=D("20"),
        ))
        quotation.items[0].container_count = D("3")
        session.flush()

        values = document_model.build_document(session, quotation).lines[0].values
        assert values["containers"] == "3"

    def test_no_capacity_means_a_blank_cell_not_a_guess(
        self, session, quotation
    ):
        """A partly-populated catalogue must not invent a figure."""
        values = document_model.build_document(session, quotation).lines[0].values
        assert values["containers"] == ""

    def test_the_seed_does_not_pin_the_column_set(self):
        """DEFAULT_COLUMNS must stay reachable in a seeded deployment.

        The seed used to store a verbatim copy of DEFAULT_COLUMNS in
        pdf_column_set. ``_column_set`` prefers a stored set, so every seeded
        deployment was frozen to the columns as they stood on install day and
        no later change to the default could reach it — the constant was dead
        code and the containers column could never have appeared.

        A stored set is meaningful only when a person chose it.
        """
        import inspect

        from seeds import seed_reference_data

        source = inspect.getsource(seed_reference_data)
        assert "pdf_column_set=None" in source, (
            "the seed must leave pdf_column_set unset so DEFAULT_COLUMNS applies"
        )

    def test_quantities_are_never_rendered_in_scientific_notation(
        self, session, quotation
    ):
        """Decimal.normalize() turns 1000.000 into 1E+3, which must not reach a
        customer document."""
        model = document_model.build_document(session, quotation)
        assert "E+" not in model.lines[0].values["quantity_packs"]

    def test_only_customer_visible_charges_are_itemised(
        self, session, sales, quotation
    ):
        quotation_service.add_charge(
            session, sales, quotation, charge_type=ChargeType.TOOLING,
            quantity=D("1"), rate=D("500"), is_customer_visible=False,
        )
        session.commit()

        model = document_model.build_document(session, quotation)
        labels = [t.label for t in model.totals]
        assert not any("Tooling" in label for label in labels)
        # But the money is still in the total the customer is asked to pay.
        assert any("Additional charges" in label for label in labels)

    def test_only_customer_visible_terms_are_included(
        self, session, sales, quotation
    ):
        term = quotation.terms[0]
        quotation_service.edit_term(
            session, sales, quotation, term.id,
            body_text=term.body_text, is_customer_visible=False,
        )
        session.commit()

        model = document_model.build_document(session, quotation)
        assert term.title not in [t.title for t in model.terms]

    def test_blank_company_fields_are_omitted_not_placeheld(
        self, session, quotation
    ):
        """With no branding configured the header simply shows less, rather than
        printing 'Address not set' at a customer."""
        model = document_model.build_document(session, quotation)
        assert model.company.address_lines == []
        assert "placeholder" not in model.company.name.lower()

    def test_percentages_are_not_padded_with_zeros(self, session, sales, quotation):
        quotation_service.update_header(
            session, sales, quotation, tax_rate_pct=D("13")
        )
        session.commit()
        model = document_model.build_document(session, quotation)
        assert any(t.label == "Tax (13%)" for t in model.totals)

    def test_a_draft_is_marked_until_it_is_approved(self, session, quotation):
        assert document_model.build_document(session, quotation).is_draft

    def test_the_filename_is_safe_and_identifies_the_revision(
        self, session, quotation
    ):
        stem = document_model.build_document(session, quotation).file_stem
        assert stem.startswith("QT-2026-0001_Rev0")
        assert all(c.isalnum() or c in "._-" for c in stem)


# --------------------------------------------------------------------------- #
# Rendering
# --------------------------------------------------------------------------- #

class TestRenderers:
    def test_pdf_renders(self, session, quotation):
        model = document_model.build_document(session, quotation)
        data = pdf_generator.render(model)
        assert data.startswith(b"%PDF-")
        assert len(data) > 1000

    def test_docx_renders(self, session, quotation):
        model = document_model.build_document(session, quotation)
        data = docx_generator.render(model)
        assert data.startswith(b"PK\x03\x04")  # a .docx is a zip
        assert len(data) > 5000

    def test_the_pdf_contains_the_quotation_content(self, session, quotation):
        from pypdf import PdfReader

        model = document_model.build_document(session, quotation)
        text = PdfReader(BytesIO(pdf_generator.render(model))).pages[0].extract_text()

        assert "QT-2026-0001" in text
        assert "Bunzl Canada" in text
        assert '12" White' in text
        assert "1,000" in text

    def test_the_docx_contains_the_quotation_content(self, session, quotation):
        from docx import Document

        model = document_model.build_document(session, quotation)
        document = Document(BytesIO(docx_generator.render(model)))

        text = "\n".join(p.text for p in document.paragraphs)
        text += "\n".join(
            cell.text for table in document.tables
            for row in table.rows for cell in row.cells
        )
        assert "QT-2026-0001" in text
        assert "Bunzl Canada" in text
        assert '12" White' in text

    def test_both_formats_agree_on_the_totals(self, session, quotation):
        from docx import Document
        from pypdf import PdfReader

        model = document_model.build_document(session, quotation)
        grand_total = next(t.amount for t in model.totals if t.emphasis)

        pdf_text = PdfReader(
            BytesIO(pdf_generator.render(model))
        ).pages[0].extract_text()
        document = Document(BytesIO(docx_generator.render(model)))
        docx_text = "\n".join(
            cell.text for table in document.tables
            for row in table.rows for cell in row.cells
        )

        assert grand_total in pdf_text
        assert grand_total in docx_text

    def test_the_draft_mark_appears_in_both_formats(self, session, quotation):
        from docx import Document
        from pypdf import PdfReader

        model = document_model.build_document(session, quotation, force_draft=True)

        pdf_text = PdfReader(
            BytesIO(pdf_generator.render(model))
        ).pages[0].extract_text()
        assert "DRAFT" in pdf_text

        document = Document(BytesIO(docx_generator.render(model)))
        assert any("D R A F T" in p.text for p in document.paragraphs)

    def test_a_final_document_carries_no_draft_mark(self, session, quotation):
        from pypdf import PdfReader

        model = document_model.build_document(session, quotation, force_draft=False)
        text = PdfReader(BytesIO(pdf_generator.render(model))).pages[0].extract_text()
        assert "DRAFT" not in text

    def test_a_long_quotation_paginates_without_error(
        self, session, sales, quotation, variant
    ):
        for _ in range(40):
            quotation_service.duplicate_line(
                session, sales, quotation, quotation.items[0].id
            )
        session.commit()

        model = document_model.build_document(session, quotation)
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(pdf_generator.render(model)))
        assert len(reader.pages) > 1
        # The table header repeats, so page two names its columns too.
        assert "Board quality" in reader.pages[1].extract_text()


class TestNoInternalDataLeaks:
    """Cost and margin must not appear in a customer document, in any format."""

    @pytest.fixture
    def costed(self, session, admin, sales, quotation, variant):
        set_cost(
            session, admin,
            CostInput(
                product_variant_id=variant.id,
                cost_per_pack=D("5.4321"),
                effective_from=JAN,
            ),
        )
        session.commit()
        # Re-add the line so it picks the cost up.
        quotation_service.remove_line(session, sales, quotation, quotation.items[0].id)
        quotation_service.add_line(
            session, sales, quotation,
            product_variant_id=variant.id,
            price_tier_code=PriceTierCode.STANDARD.value,
            quantity_packs=D("1000"),
        )
        quotation_service.update_line(
            session, sales, quotation, quotation.items[0].id,
            internal_remarks="MARGIN IS THIN — do not disclose",
        )
        session.commit()
        assert quotation.total_cost == D("5432.10")
        return quotation

    def test_the_model_has_no_cost_fields_at_all(self, session, costed):
        model = document_model.build_document(session, costed)
        blob = repr(model)
        assert "5432.10" not in blob
        assert "5.4321" not in blob
        assert "do not disclose" not in blob

    def test_cost_does_not_appear_in_the_pdf_bytes(self, session, costed):
        model = document_model.build_document(session, costed)
        from pypdf import PdfReader

        text = PdfReader(BytesIO(pdf_generator.render(model))).pages[0].extract_text()
        assert "5,432.10" not in text
        assert "5.4321" not in text
        assert "do not disclose" not in text

    def test_cost_does_not_appear_in_the_docx(self, session, costed):
        from docx import Document

        model = document_model.build_document(session, costed)
        document = Document(BytesIO(docx_generator.render(model)))
        text = "\n".join(p.text for p in document.paragraphs) + "\n".join(
            cell.text for table in document.tables
            for row in table.rows for cell in row.cells
        )
        assert "5,432.10" not in text
        assert "5.4321" not in text
        assert "do not disclose" not in text

    def test_internal_only_charges_are_not_named(self, session, sales, costed):
        quotation_service.add_charge(
            session, sales, costed, charge_type=ChargeType.TOOLING,
            description="Secret tooling deal", quantity=D("1"), rate=D("500"),
            is_customer_visible=False,
        )
        session.commit()

        model = document_model.build_document(session, costed)
        assert "Secret tooling deal" not in repr(model)


# --------------------------------------------------------------------------- #
# The release gate
# --------------------------------------------------------------------------- #

class TestReleaseGate:
    def test_a_draft_cannot_be_released(self, session, quotation):
        blockers = approval_service.release_blockers(session, quotation)
        assert blockers
        assert not document_service.can_release(session, quotation)

    def test_a_draft_copy_is_still_available(self, session, sales, quotation):
        generated = document_service.generate(
            session, sales, quotation, DocumentFormat.PDF
        )
        assert generated.is_draft
        assert generated.filename.endswith("_DRAFT.pdf")

    @pytest.mark.parametrize("fmt", list(DocumentFormat))
    def test_a_final_document_is_refused_while_blocked(
        self, session, sales, quotation, fmt
    ):
        """Neither format escapes the gate."""
        with pytest.raises(ApprovalError, match="cannot be produced yet"):
            document_service.generate(session, sales, quotation, fmt, draft=False)

    def test_release_is_allowed_once_approved(
        self, session, sales, manager, quotation
    ):
        _approve_and_issue(session, quotation, sales, manager)
        assert document_service.can_release(session, quotation)

    @pytest.mark.parametrize("fmt", list(DocumentFormat))
    def test_both_formats_can_be_produced_after_approval(
        self, session, sales, manager, quotation, fmt
    ):
        _approve_and_issue(session, quotation, sales, manager)
        generated = document_service.generate(session, sales, quotation, fmt, draft=False)
        session.commit()
        assert not generated.is_draft
        assert not generated.filename.endswith("_DRAFT" + generated.filename[-5:])

    def test_a_final_document_is_archived_with_its_hash(
        self, session, sales, manager, quotation
    ):
        _approve_and_issue(session, quotation, sales, manager)
        document_service.generate(
            session, sales, quotation, DocumentFormat.PDF, draft=False
        )
        session.commit()

        stored = document_service.stored_documents(session, quotation.id)
        assert len(stored) == 1
        assert stored[0].sha256 and len(stored[0].sha256) == 64

    def test_a_draft_is_not_archived(self, session, sales, quotation):
        document_service.generate(session, sales, quotation, DocumentFormat.PDF)
        session.commit()
        assert document_service.stored_documents(session, quotation.id) == []

    def test_an_archived_document_can_be_re_downloaded(
        self, session, sales, manager, quotation
    ):
        _approve_and_issue(session, quotation, sales, manager)
        original = document_service.generate(
            session, sales, quotation, DocumentFormat.DOCX, draft=False
        )
        session.commit()

        stored = document_service.stored_documents(session, quotation.id)
        again = document_service.fetch(session, sales, stored[0].id)
        assert again.data == original.data

    def test_generation_is_audited(self, session, sales, quotation):
        document_service.generate(session, sales, quotation, DocumentFormat.DOCX)
        session.commit()
        entry = (
            session.query(AuditLog)
            .filter_by(action=AuditAction.PDF_GENERATED.value)
            .one()
        )
        assert entry.new_value_json["format"] == "DOCX"
        assert entry.new_value_json["draft"] is True

    def test_a_user_without_permission_cannot_generate(
        self, session, make_auth_user, quotation
    ):
        pricer = make_auth_user(RoleCode.PRICING_ADMIN.value)
        with pytest.raises(PermissionDenied):
            document_service.generate(session, pricer, quotation, DocumentFormat.PDF)


# --------------------------------------------------------------------------- #
# Approvals
# --------------------------------------------------------------------------- #

class TestApproval:
    def test_within_authority_goes_straight_to_approved(
        self, session, admin, quotation
    ):
        """A System Administrator has no limits, so nothing is triggered."""
        approval = approval_service.submit(session, quotation, admin)
        session.commit()
        assert approval is None
        assert quotation.status is QuotationStatus.APPROVED

    def test_a_value_above_authority_raises_a_request(
        self, session, sales, quotation
    ):
        # Sales authority is 25,000; this quotation is 8,220.
        quotation_service.update_line(
            session, sales, quotation, quotation.items[0].id,
            quantity_packs=D("10000"),
        )
        session.commit()

        approval = approval_service.submit(session, quotation, sales)
        session.commit()
        assert approval is not None
        assert quotation.status is QuotationStatus.PENDING_APPROVAL
        triggers = [t["trigger"] for t in approval.triggered_rules_json]
        assert "VALUE_ABOVE_AUTHORITY" in triggers

    def test_a_discount_above_the_limit_raises_a_request(
        self, session, sales, quotation
    ):
        quotation_service.update_header(
            session, sales, quotation, quote_discount_pct=D("12")
        )
        session.commit()
        approval = approval_service.submit(session, quotation, sales)
        session.commit()
        assert "DISCOUNT_ABOVE_LIMIT" in [
            t["trigger"] for t in approval.triggered_rules_json
        ]

    def test_an_incomplete_quotation_cannot_be_submitted(
        self, session, sales, quotation
    ):
        quotation_service.remove_line(session, sales, quotation, quotation.items[0].id)
        session.commit()
        with pytest.raises(ApprovalError, match="not complete"):
            approval_service.submit(session, quotation, sales)

    def test_a_manager_can_approve_someone_elses(
        self, session, sales, manager, quotation
    ):
        quotation_service.update_header(
            session, sales, quotation, quote_discount_pct=D("12")
        )
        session.commit()
        approval = approval_service.submit(session, quotation, sales)
        session.commit()

        approval_service.approve(session, quotation, manager, approval.id, "fine")
        session.commit()
        assert quotation.status is QuotationStatus.APPROVED

    def test_the_submitter_cannot_approve_their_own(
        self, session, manager, quotation
    ):
        """A custom price triggers approval for every role, so this exercises
        the self-approval block rather than skipping when nothing fires."""
        quotation.sales_user_id = manager.id
        quotation_service.change_line_tier(
            session, manager, quotation, quotation.items[0].id,
            PriceTierCode.CUSTOM.value,
            custom_price_per_pack=D("7.00"), custom_price_reason="volume",
        )
        session.commit()

        approval = approval_service.submit(session, quotation, manager)
        session.commit()
        assert approval is not None, "a custom price must require approval"

        with pytest.raises(ApprovalError, match="cannot also approve"):
            approval_service.approve(session, quotation, manager, approval.id)

    def test_a_system_administrator_cannot_approve_their_own_either(
        self, session, sales, admin, quotation
    ):
        quotation_service.update_header(
            session, sales, quotation, quote_discount_pct=D("12")
        )
        session.commit()
        approval = approval_service.submit(session, quotation, sales)
        session.commit()

        # The admin owns nothing here, but make them the salesperson.
        quotation.sales_user_id = admin.id
        session.commit()
        with pytest.raises(ApprovalError, match="your own quotation"):
            approval_service.approve(session, quotation, admin, approval.id)

    def test_the_queue_hides_your_own_work(self, session, sales, manager, quotation):
        quotation_service.update_header(
            session, sales, quotation, quote_discount_pct=D("12")
        )
        session.commit()
        approval_service.submit(session, quotation, sales)
        session.commit()

        assert len(approval_service.queue(session, manager)) == 1
        quotation.sales_user_id = manager.id
        session.commit()
        assert approval_service.queue(session, manager) == []

    def test_rejecting_requires_a_reason(self, session, sales, manager, quotation):
        quotation_service.update_header(
            session, sales, quotation, quote_discount_pct=D("12")
        )
        session.commit()
        approval = approval_service.submit(session, quotation, sales)
        session.commit()

        with pytest.raises(ApprovalError, match="reason is required"):
            approval_service.reject(session, quotation, manager, approval.id, "  ")

    def test_rejecting_moves_the_quotation_and_records_the_reason(
        self, session, sales, manager, quotation
    ):
        quotation_service.update_header(
            session, sales, quotation, quote_discount_pct=D("12")
        )
        session.commit()
        approval = approval_service.submit(session, quotation, sales)
        session.commit()

        approval_service.reject(
            session, quotation, manager, approval.id, "margin too thin"
        )
        session.commit()
        assert quotation.status is QuotationStatus.REJECTED_INTERNALLY
        assert approval.rejection_reason == "margin too thin"

    def test_returning_for_revision(self, session, sales, manager, quotation):
        quotation_service.update_header(
            session, sales, quotation, quote_discount_pct=D("12")
        )
        session.commit()
        approval = approval_service.submit(session, quotation, sales)
        session.commit()

        approval_service.return_for_revision(
            session, quotation, manager, approval.id, "add the 14 inch"
        )
        session.commit()
        assert quotation.status is QuotationStatus.REVISION_REQUIRED

    def test_a_decided_request_cannot_be_decided_again(
        self, session, sales, manager, quotation
    ):
        quotation_service.update_header(
            session, sales, quotation, quote_discount_pct=D("12")
        )
        session.commit()
        approval = approval_service.submit(session, quotation, sales)
        session.commit()
        approval_service.approve(session, quotation, manager, approval.id)
        session.commit()

        with pytest.raises(ApprovalError, match="already"):
            approval_service.approve(session, quotation, manager, approval.id)

    def test_approving_past_a_blocking_warning_needs_an_override_reason(
        self, session, sales, manager, quotation, variant
    ):
        quotation_service.change_line_tier(
            session, sales, quotation, quotation.items[0].id,
            PriceTierCode.CUSTOM.value,
            custom_price_per_pack=D("1.00"), custom_price_reason="strategic",
        )
        session.commit()
        approval = approval_service.submit(session, quotation, sales)
        session.commit()

        with pytest.raises(ApprovalError, match="override reason"):
            approval_service.approve(session, quotation, manager, approval.id)

        approval_service.approve(
            session, quotation, manager, approval.id,
            override_reason="board approved this account",
        )
        session.commit()
        assert quotation.status is QuotationStatus.APPROVED

        entry = (
            session.query(AuditLog)
            .filter_by(action=AuditAction.WARNING_OVERRIDDEN.value)
            .one()
        )
        assert entry.reason == "board approved this account"


# --------------------------------------------------------------------------- #
# Revisions
# --------------------------------------------------------------------------- #

class TestRevisions:
    def _issue(self, session, sales, manager, quotation):
        _approve_and_issue(session, quotation, sales, manager)
        document_service.generate(
            session, sales, quotation, DocumentFormat.PDF, draft=False
        )
        stored = document_service.stored_documents(session, quotation.id)
        record = revision_service.issue(
            session, sales, quotation, pdf_attachment_id=stored[0].id
        )
        session.commit()
        return record

    def test_issuing_locks_and_snapshots(self, session, sales, manager, quotation):
        record = self._issue(session, sales, manager, quotation)
        assert quotation.is_locked
        assert quotation.issued_at is not None
        assert record.snapshot_json["quote_number"] == "QT-2026-0001"
        assert len(record.snapshot_json["lines"]) == 1
        assert record.new_pdf_attachment_id is not None

    def test_a_snapshot_carries_no_cost(self, session, admin, sales, manager,
                                        quotation, variant):
        """A snapshot may be shown to anyone who can view the quotation, and
        cost visibility is a separate permission."""
        set_cost(
            session, admin,
            CostInput(
                product_variant_id=variant.id, cost_per_pack=D("5.4321"),
                effective_from=JAN,
            ),
        )
        session.commit()
        record = self._issue(session, sales, manager, quotation)

        for line in record.snapshot_json["lines"]:
            assert not any("cost" in key for key in line)
        assert "5.4321" not in str(record.snapshot_json)

    def test_issuing_twice_is_refused(self, session, sales, manager, quotation):
        self._issue(session, sales, manager, quotation)
        with pytest.raises(RevisionError, match="already been issued"):
            revision_service.issue(session, sales, quotation)

    def test_an_issued_quotation_cannot_be_edited(
        self, session, sales, manager, quotation
    ):
        self._issue(session, sales, manager, quotation)
        with pytest.raises(PermissionDenied, match="Create a revision"):
            quotation_service.update_header(session, sales, quotation, brand="X")

    def test_a_revision_needs_a_reason(self, session, sales, manager, quotation):
        self._issue(session, sales, manager, quotation)
        with pytest.raises(RevisionError, match="reason is required"):
            revision_service.create_revision(session, sales, quotation, "  ")

    def test_an_unissued_quotation_does_not_need_a_revision(
        self, session, sales, quotation
    ):
        with pytest.raises(RevisionError, match="can be edited directly"):
            revision_service.create_revision(session, sales, quotation, "why")

    def test_a_revision_copies_everything_and_supersedes(
        self, session, sales, manager, quotation
    ):
        self._issue(session, sales, manager, quotation)
        revised = revision_service.create_revision(
            session, sales, quotation, "customer added a size"
        )
        session.commit()

        assert revised.revision_no == 1
        assert revised.quote_number == quotation.quote_number
        assert revised.is_current_revision
        assert not revised.is_locked
        assert revised.status is QuotationStatus.DRAFT
        assert len(revised.items) == len(quotation.items)
        assert len(revised.charges) == len(quotation.charges)
        assert len(revised.terms) == len(quotation.terms)

        assert not quotation.is_current_revision
        assert quotation.is_locked

    def test_editing_a_revision_leaves_the_original_untouched(
        self, session, sales, manager, quotation
    ):
        self._issue(session, sales, manager, quotation)
        original_total = quotation.grand_total

        revised = revision_service.create_revision(session, sales, quotation, "more")
        session.commit()
        quotation_service.update_line(
            session, sales, revised, revised.items[0].id, quantity_packs=D("2000")
        )
        session.commit()

        assert revised.grand_total != original_total
        assert session.get(Quotation, quotation.id).grand_total == original_total

    def test_the_diff_reports_exactly_what_changed(
        self, session, sales, manager, quotation
    ):
        self._issue(session, sales, manager, quotation)
        before = revision_service.snapshot(quotation)

        revised = revision_service.create_revision(session, sales, quotation, "more")
        session.commit()
        quotation_service.update_line(
            session, sales, revised, revised.items[0].id, quantity_packs=D("2000")
        )
        quotation_service.update_header(session, sales, revised, brand="Bunzl Foodservice")
        session.commit()

        diff = revision_service.compare(before, revision_service.snapshot(revised))
        assert revision_service.has_changes(diff)
        assert any(c["field"] == "brand" for c in diff["header"])
        assert any(c["field"] == "grand_total" for c in diff["totals"])

        changed = [line for line in diff["lines"] if line["change"] == "changed"]
        assert len(changed) == 1
        assert "quantity_packs" in changed[0]["fields"]

    def test_an_unchanged_revision_reports_no_differences(
        self, session, sales, manager, quotation
    ):
        self._issue(session, sales, manager, quotation)
        before = revision_service.snapshot(quotation)
        revised = revision_service.create_revision(session, sales, quotation, "reprint")
        session.commit()

        diff = revision_service.compare(before, revision_service.snapshot(revised))
        assert diff["lines"] == []
        assert diff["totals"] == []

    def test_both_revisions_remain_listed(self, session, sales, manager, quotation):
        self._issue(session, sales, manager, quotation)
        revision_service.create_revision(session, sales, quotation, "more")
        session.commit()

        family = revision_service.revisions_for(session, quotation.root_quotation_id)
        assert [q.revision_no for q in family] == [0, 1]

    def test_a_revision_snapshot_cannot_be_tampered_with(
        self, session, sales, manager, quotation
    ):
        record = self._issue(session, sales, manager, quotation)
        record.change_reason = "tampering"
        with pytest.raises(ImmutableRecordError):
            session.commit()
        session.rollback()


class TestOnePageAndOrdering:
    """Two changes to how the quotation is laid out, both customer-visible."""

    def test_a_short_quotation_comes_out_on_one_page(
        self, session, sales, quotation
    ):
        """It spilled a signature block onto a second sheet.

        A quotation that runs a few centimetres over is a worse document than
        the same one set tighter: the reader turns the page for a signature
        line and nothing else.
        """
        from pypdf import PdfReader

        model = document_model.build_document(session, quotation)
        pdf = pdf_generator.render(model)
        assert len(PdfReader(BytesIO(pdf)).pages) == 1

    def test_a_long_quotation_is_not_shrunk_to_no_purpose(
        self, session, sales, quotation
    ):
        """Only a retry that reaches one page is used.

        Forty lines were never going to fit, so shrinking the type would cost
        legibility and buy nothing. The full-size render is returned instead.
        """
        for _ in range(40):
            quotation_service.duplicate_line(
                session, sales, quotation, quotation.items[0].id
            )
        session.commit()

        from pypdf import PdfReader

        model = document_model.build_document(session, quotation)
        returned = len(PdfReader(BytesIO(pdf_generator.render(model))).pages)
        _, at_full = pdf_generator._render_at(model, "A4", 1.0)
        _, at_smallest = pdf_generator._render_at(model, "A4", 0.70)

        # Compared by page count, not bytes: a PDF embeds a creation timestamp,
        # so two renders of identical content are never byte-identical.
        assert at_smallest < at_full, (
            "the fixture is too small to distinguish the two renders"
        )
        assert returned == at_full, "a multi-page quotation was needlessly shrunk"

    def test_the_shipping_terms_print_below_the_conditions(
        self, session, sales, admin, quotation, carrier
    ):
        """Incoterms are conditions of sale, not a footnote to the total.

        They sat between the total and the notes, separating the figure from
        everything explaining it.
        """
        from pypdf import PdfReader

        from modules import shipping_service
        from modules.constants import ContainerSize

        shipping_service.add_container(
            session, admin, quotation, shipping_line_id=carrier.id,
            container_size=ContainerSize.FORTY_FT_HC,
            container_count=D("1"), freight_cost=D("4400"),
        )
        session.commit()

        model = document_model.build_document(session, quotation)
        text = "\n".join(
            p.extract_text() or ""
            for p in PdfReader(BytesIO(pdf_generator.render(model))).pages
        )
        terms_at = text.find("Terms and conditions")
        incoterms_at = text.find("Incoterms")
        assert terms_at != -1 and incoterms_at != -1
        assert terms_at < incoterms_at, "the shipping summary is still above the terms"

    def test_the_word_document_orders_them_the_same_way(
        self, session, sales, admin, quotation, carrier
    ):
        """One model, two renderers; they must not disagree about order."""
        from docx import Document as DocxDocument

        from modules import shipping_service
        from modules.constants import ContainerSize

        shipping_service.add_container(
            session, admin, quotation, shipping_line_id=carrier.id,
            container_size=ContainerSize.FORTY_FT_HC,
            container_count=D("1"), freight_cost=D("4400"),
        )
        session.commit()

        model = document_model.build_document(session, quotation)
        doc = DocxDocument(BytesIO(docx_generator.render(model)))
        body = [p.text for p in doc.paragraphs]
        terms_at = next(i for i, t in enumerate(body) if "Terms and conditions" in t)
        incoterms_at = next(i for i, t in enumerate(body) if "Incoterms" in t)
        assert terms_at < incoterms_at
