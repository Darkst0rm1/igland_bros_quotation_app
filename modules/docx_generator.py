"""Word renderer, python-docx.

Consumes the same :class:`~modules.document_model.QuotationDocument` as
:mod:`modules.pdf_generator`, so the two formats cannot say different things
about the same quotation number.

Two behaviours need raw OOXML because python-docx has no API for them: marking
a table row as a repeating header, and the page-number field in the footer.
Both are small, well-defined bits of the format and are isolated in helpers at
the top of this module.

Note for anyone relying on these files commercially: a ``.docx`` is editable by
whoever receives it, so it is not evidence of what was sent. The PDF is the
record — see ``revision_service``, which stores the PDF as the revision
artefact.
"""

from __future__ import annotations

import logging
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from modules.document_model import QuotationDocument
from modules.utilities import format_date

log = logging.getLogger(__name__)

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5F, 0x6B, 0x7A)
DRAFT_RED = RGBColor(0xC0, 0x39, 0x2B)
BAND = "EEF1F5"


# --------------------------------------------------------------------------- #
# Raw OOXML helpers
# --------------------------------------------------------------------------- #

def _repeat_header(row) -> None:  # noqa: ANN001
    """Mark a table row as a header that repeats on every page.

    python-docx exposes no API for ``tblHeader``; without it a table spanning
    pages loses its column headings on page two, which is the single most
    common complaint about generated Word tables.
    """
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _keep_row_together(row) -> None:  # noqa: ANN001
    """Stop a single row from splitting across a page break."""
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def _shade(cell, hex_colour: str) -> None:  # noqa: ANN001
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shading)


def _page_number_field(paragraph) -> None:  # noqa: ANN001
    """Insert 'Page N of M' as a live field.

    Written as field codes rather than literal text so the numbers stay correct
    if the recipient edits the document — which, this being Word, they will.
    """
    def field(instruction: str) -> None:
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run = paragraph.add_run()
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    paragraph.add_run("Page ")
    field("PAGE")
    paragraph.add_run(" of ")
    field("NUMPAGES")


# --------------------------------------------------------------------------- #
# Small builders
# --------------------------------------------------------------------------- #

def _para(container, text: str, *, size=9, bold=False, colour=INK, align=None,  # noqa: ANN001
          space_after=2):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(0)
    if align is not None:
        paragraph.alignment = align
    for index, line in enumerate(str(text).split("\n")):
        run = paragraph.add_run(line)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = colour
        if index < len(str(text).split("\n")) - 1:
            run.add_break()
    return paragraph


def _label(container, text: str):  # noqa: ANN001
    return _para(container, text.upper(), size=7, colour=MUTED, space_after=0)


# --------------------------------------------------------------------------- #
# Sections
# --------------------------------------------------------------------------- #

def _write_header(doc: Document, model: QuotationDocument) -> None:
    if model.is_draft:
        # A true diagonal watermark in Word needs a VML shape in the header and
        # is fragile across versions. A full-width banner is unmissable, edits
        # cleanly, and survives conversion.
        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = banner.add_run("D R A F T  —  N O T  A  F I R M  O F F E R")
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = DRAFT_RED

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells

    if model.company.logo_bytes:
        try:
            left.paragraphs[0].add_run().add_picture(
                BytesIO(model.company.logo_bytes), width=Inches(1.8)
            )
        except Exception:  # noqa: BLE001 - a bad logo must not stop a quotation
            log.warning("Logo could not be embedded; continuing without it")

    _para(left, model.company.name, size=15, bold=True, space_after=1)
    for line in model.company.address_lines:
        _para(left, line, size=8, colour=MUTED, space_after=0)
    if model.company.contact_line:
        _para(left, model.company.contact_line, size=8, colour=MUTED, space_after=0)
    if model.company.tax_number:
        _para(left, f"Tax no. {model.company.tax_number}", size=8, colour=MUTED,
              space_after=0)

    _para(right, model.title, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    meta = [
        ("Quotation no.", f"{model.quote_number}  {model.revision_label}"),
        ("Quote date", format_date(model.quote_date)),
    ]
    if model.valid_until:
        meta.append(("Valid until", format_date(model.valid_until)))
    for label, value in meta:
        paragraph = right.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = MUTED
        value_run = paragraph.add_run(value)
        value_run.font.size = Pt(9)
        value_run.font.bold = True

    _para(doc, "", size=4)


def _write_customer(doc: Document, model: QuotationDocument) -> None:
    contact_bits = [
        b for b in (
            model.customer.contact_name,
            model.customer.contact_email,
            model.customer.contact_phone,
        ) if b
    ]
    blocks = [("Prepared for", model.customer.company), ("Contact", "\n".join(contact_bits))]
    if model.customer.project:
        blocks.append(("Project", model.customer.project))
    if model.customer.brand:
        blocks.append(("Brand", model.customer.brand))
    if model.customer.distributor:
        blocks.append(("Distributor", model.customer.distributor))
    if model.customer.purchase_order:
        blocks.append(("PO reference", model.customer.purchase_order))
    if model.customer.billing_address:
        blocks.append(("Billing address", model.customer.billing_address))
    if model.customer.shipping_address:
        blocks.append(("Shipping address", model.customer.shipping_address))

    for start in range(0, len(blocks), 3):
        chunk = blocks[start:start + 3]
        table = doc.add_table(rows=1, cols=3)
        for index, cell in enumerate(table.rows[0].cells):
            if index < len(chunk):
                label, value = chunk[index]
                _label(cell, label)
                _para(cell, value or "—", size=9)
    _para(doc, "", size=4)


def _write_lines(doc: Document, model: QuotationDocument) -> None:
    numeric = set(model.numeric_column_indexes)
    table = doc.add_table(rows=1, cols=len(model.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row = table.rows[0]
    _repeat_header(header_row)
    for index, heading in enumerate(model.column_headings):
        cell = header_row.cells[index]
        _shade(cell, BAND)
        _para(
            cell, heading, size=7.5, bold=True,
            align=WD_ALIGN_PARAGRAPH.RIGHT if index in numeric else None,
            space_after=0,
        )

    for line in model.lines:
        row = table.add_row()
        _keep_row_together(row)
        for index, value in enumerate(line.cells(model.columns)):
            _para(
                row.cells[index], value, size=8,
                align=WD_ALIGN_PARAGRAPH.RIGHT if index in numeric else None,
                space_after=0,
            )
    _para(doc, "", size=4)


def _write_totals(doc: Document, model: QuotationDocument) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    for total in model.totals:
        row = table.add_row()
        _keep_row_together(row)
        # The grand total gets the banded, large-type treatment the PDF gives
        # it. Both renderers read from the same model, so letting them diverge
        # visually would mean the customer sees a different document depending
        # on which format the employee happened to download.
        _para(
            row.cells[0],
            total.label.upper() if total.emphasis else total.label,
            size=9, bold=total.emphasis,
            align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0,
        )
        _para(
            row.cells[1], total.amount,
            size=14 if total.emphasis else 9, bold=total.emphasis,
            align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0,
        )
        if total.emphasis:
            for cell in row.cells:
                _shade(cell, "EEF1F5")
    _para(doc, "", size=4)


def _write_shipping(doc: Document, model: QuotationDocument) -> None:
    """The shipping summary. No per-container table — see ``pdf_generator``.

    Both renderers consume the same model, so the table disappearing from one
    and not the other is not possible: ``DocumentShipping`` no longer carries
    rows for either to print.
    """
    shipping = model.shipping
    if not shipping:
        return

    if shipping.summary:
        _para(
            doc,
            "  ·  ".join(f"{label}: {value}" for label, value in shipping.summary),
            size=8.5,
        )
    if shipping.freight_statement:
        _para(doc, shipping.freight_statement, size=8.5)
    if shipping.notes:
        _para(doc, shipping.notes, size=8.5)


def _write_terms(doc: Document, model: QuotationDocument) -> None:
    if model.customer_notes:
        _para(doc, "Notes", size=11, bold=True, space_after=4)
        _para(doc, model.customer_notes, size=8.5)

    if not model.terms:
        return
    _para(doc, "Terms and conditions", size=11, bold=True, space_after=4)
    for term in model.terms:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        # keep_with_next stops a heading being stranded at the foot of a page.
        paragraph.paragraph_format.keep_with_next = True
        title_run = paragraph.add_run(f"{term.title}: ")
        title_run.font.size = Pt(8.5)
        title_run.font.bold = True
        body_run = paragraph.add_run(term.body)
        body_run.font.size = Pt(8.5)


def _write_signatures(doc: Document, model: QuotationDocument) -> None:
    _para(doc, "", size=6)
    table = doc.add_table(rows=1, cols=3)
    prepared, approved, dated = table.rows[0].cells

    _label(prepared, "Prepared by")
    _para(prepared, model.prepared_by or model.signature_name or "—", size=9)
    if model.prepared_by_title or model.signature_title:
        _para(prepared, model.prepared_by_title or model.signature_title,
              size=8, colour=MUTED)

    _label(approved, "Approved by")
    _para(approved, model.approved_by or "—", size=9)

    _label(dated, "Date")
    _para(dated, format_date(model.quote_date), size=9)

    if model.show_acceptance_line:
        _para(doc, "", size=8)
        _label(doc, "Customer acceptance")
        acceptance = doc.add_table(rows=1, cols=3)
        acceptance.style = "Table Grid"
        for cell, label in zip(
            acceptance.rows[0].cells, ("Signature", "Name", "Date"), strict=False
        ):
            _para(cell, "", size=9)
            _para(cell, label, size=8, colour=MUTED, space_after=0)


def _write_footer(doc: Document, model: QuotationDocument) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)

    bits = [model.company.name, model.company.contact_line]
    lead = "  ·  ".join(b for b in bits if b)
    if lead:
        run = paragraph.add_run(lead + "  ·  ")
        run.font.size = Pt(7)
        run.font.color.rgb = MUTED

    reference = paragraph.add_run(f"{model.quote_number} {model.revision_label}  ·  ")
    reference.font.size = Pt(7)
    reference.font.color.rgb = MUTED
    _page_number_field(paragraph)
    for run in paragraph.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = MUTED

    if model.confidentiality_text:
        note = footer.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run(model.confidentiality_text)
        run.font.size = Pt(6.5)
        run.font.color.rgb = MUTED


# --------------------------------------------------------------------------- #
# Entry point
# --------------------------------------------------------------------------- #

def render(document: QuotationDocument, page_size: str = "A4") -> bytes:
    """Render the document and return the .docx bytes."""
    doc = Document()

    section = doc.sections[0]
    if (page_size or "A4").upper() == "A4":
        section.page_width, section.page_height = Inches(8.27), Inches(11.69)
    else:
        section.page_width, section.page_height = Inches(8.5), Inches(11)
    if len(document.columns) > 8:
        # Match the PDF: a wide column set is rotated rather than truncated.
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    for margin in ("left_margin", "right_margin"):
        setattr(section, margin, Inches(0.6))
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    _write_header(doc, document)
    _write_customer(doc, document)
    _write_lines(doc, document)
    _write_totals(doc, document)
    _write_shipping(doc, document)
    _write_terms(doc, document)
    _write_signatures(doc, document)

    if document.thank_you_text:
        _para(doc, "", size=6)
        _para(doc, document.thank_you_text, size=8, colour=MUTED,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    _write_footer(doc, document)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
